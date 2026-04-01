#!/usr/bin/env python3
"""Build professionally styled Word docs for Billy's resumes."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import copy

NAVY = RGBColor(0x0A, 0x26, 0x47)
GRAY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
BODY = RGBColor(0x2D, 0x2D, 0x2D)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = __import__('lxml.etree', fromlist=['SubElement']).SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            element = __import__('lxml.etree', fromlist=['SubElement']).SubElement(tcBorders, qn(f'w:{edge}'))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), val.get('space', '0'))

def make_no_border(cell):
    set_cell_border(cell,
        top={"val": "none", "sz": "0", "color": "FFFFFF"},
        bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
        start={"val": "none", "sz": "0", "color": "FFFFFF"},
        end={"val": "none", "sz": "0", "color": "FFFFFF"})

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    p.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = __import__('lxml.etree', fromlist=['SubElement']).SubElement(pPr, qn('w:pBdr'))
    bottom = __import__('lxml.etree', fromlist=['SubElement']).SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'D0D8E0')

def add_entry_header(doc, title, date):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(2.1)

    # Title cell
    cell_l = table.cell(0, 0)
    cell_l.paragraphs[0].space_before = Pt(4)
    cell_l.paragraphs[0].space_after = Pt(0)
    run = cell_l.paragraphs[0].add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLACK
    run.font.name = 'Calibri'
    make_no_border(cell_l)

    # Date cell
    cell_r = table.cell(0, 1)
    cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell_r.paragraphs[0].space_before = Pt(4)
    cell_r.paragraphs[0].space_after = Pt(0)
    run = cell_r.paragraphs[0].add_run(date)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'
    make_no_border(cell_r)

def add_org_line(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = Pt(1)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.space_before = Pt(0)
    p.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = BODY
    run.font.name = 'Calibri'

def add_cert_bullet(doc, name, detail):
    p = doc.add_paragraph(style='List Bullet')
    p.space_before = Pt(0)
    p.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    run.font.name = 'Calibri'
    run = p.add_run(f' - {detail}')
    run.font.size = Pt(9.5)
    run.font.color.rgb = BODY
    run.font.name = 'Calibri'

def add_edu_row(doc, name, detail):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(4.0)
    table.columns[1].width = Inches(3.1)

    cell_l = table.cell(0, 0)
    cell_l.paragraphs[0].space_before = Pt(1)
    cell_l.paragraphs[0].space_after = Pt(1)
    run = cell_l.paragraphs[0].add_run(name)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    run.font.name = 'Calibri'
    make_no_border(cell_l)

    cell_r = table.cell(0, 1)
    cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell_r.paragraphs[0].space_before = Pt(1)
    cell_r.paragraphs[0].space_after = Pt(1)
    run = cell_r.paragraphs[0].add_run(detail)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'
    make_no_border(cell_r)

def add_skill_grid(doc, skills):
    """Add skills as a 2-column table of bullets."""
    rows = (len(skills) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(3.55)
    table.columns[1].width = Inches(3.55)

    for i, skill in enumerate(skills):
        row_idx = i // 2
        col_idx = i % 2
        cell = table.cell(row_idx, col_idx)
        cell.paragraphs[0].space_before = Pt(0)
        cell.paragraphs[0].space_after = Pt(0)
        run = cell.paragraphs[0].add_run(f'  \u2022  {skill}')
        run.font.size = Pt(9.5)
        run.font.color.rgb = BODY
        run.font.name = 'Calibri'
        make_no_border(cell)

    # Clear any empty cells
    if len(skills) % 2 == 1:
        cell = table.cell(rows - 1, 1)
        make_no_border(cell)

def build_header(doc, name, tagline, location, phone, email):
    # Name + tagline on left, contact on right
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(5.0)
    table.columns[1].width = Inches(2.1)

    # Left cell - name and tagline
    cell_l = table.cell(0, 0)
    p = cell_l.paragraphs[0]
    p.space_after = Pt(0)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY
    run.font.name = 'Calibri'

    p2 = cell_l.add_paragraph()
    p2.space_before = Pt(0)
    p2.space_after = Pt(0)
    run = p2.add_run(tagline)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4A, 0x7F, 0xB5)
    run.font.name = 'Calibri'
    make_no_border(cell_l)

    # Right cell - contact
    cell_r = table.cell(0, 1)
    cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cell_r.paragraphs[0].space_after = Pt(0)
    run = cell_r.paragraphs[0].add_run(location)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

    for text in [phone, email]:
        p = cell_r.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.space_before = Pt(0)
        p.space_after = Pt(0)
        run = p.add_run(text)
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.font.name = 'Calibri'
    make_no_border(cell_r)

    # Divider line
    p = doc.add_paragraph()
    p.space_before = Pt(4)
    p.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = __import__('lxml.etree', fromlist=['SubElement']).SubElement(pPr, qn('w:pBdr'))
    bottom = __import__('lxml.etree', fromlist=['SubElement']).SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0A2647')

def add_summary(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(0)
    p.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x3A, 0x3A, 0x3A)
    run.font.name = 'Calibri'

def setup_doc():
    doc = Document()
    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = BODY
    # Tighten paragraph spacing
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    return doc


# ═══════════════════════════════════════════════
# BUILD PCA RESUME
# ═══════════════════════════════════════════════
doc = setup_doc()

build_header(doc,
    "Billy James Cruz",
    "Patient Care Assistant (PCA) Candidate | EMT-Certified | BLS Provider (AHA)",
    "Houston, TX Area", "(409) 673-4859", "billyjcruz@icloud.com")

add_section_heading(doc, "Professional Summary")
add_summary(doc, "EMT-certified Patient Care Assistant candidate with direct patient care experience on ambulance clinical rotations with the Seguin Fire Department. Trained in patient assessment, vital signs, ADL support (hygiene, nutrition, ambulation, elimination), and HIPAA-compliant documentation under licensed supervision. Two years leading a team of 8 at a 2,000-member facility. Conversational Spanish. Available for day and night 12-hour shifts. Committed to Houston Methodist's ICARE values.")

add_section_heading(doc, "Skills")
add_skill_grid(doc, [
    "Direct Patient Care", "Basic Life Support (BLS/AHA)",
    "Patient Assessment and Vital Signs", "Activities of Daily Living (ADLs)",
    "Patient Documentation and Charting", "HIPAA Compliance",
    "Fall and Pressure Injury Prevention", "Infection Prevention and Control",
    "Hourly Rounding and Four P's", "Bedside Shift Handoff",
    "Data Entry, Collection, and Interpretation", "Conversational Spanish",
])

add_section_heading(doc, "Healthcare Training and Clinical Experience")

add_entry_header(doc, "Emergency Medical Technician", "Jan 2025 - May 2025")
add_org_line(doc, "EMS University | San Antonio, TX | Clinical Rotations: Seguin Fire Department")
add_bullet(doc, "Provided direct patient care on emergency ambulance calls during Seguin Fire Department clinical rotations under licensed supervision.")
add_bullet(doc, "Obtained vital signs (BP, pulse, SpO2, respirations), monitored condition changes, and documented findings per HIPAA protocol.")
add_bullet(doc, "Provided ADL support: patient hygiene, comfort positioning, ambulation aid, and nutrition monitoring during stabilization.")
add_bullet(doc, "Assisted with oxygen delivery, medication administration, and c-spine stabilization; delivered verbal handoffs to ER staff.")

add_entry_header(doc, "Recruit Fire Academy", "Aug 2024 - Oct 2024")
add_org_line(doc, "TEEX | College Station, TX")
add_bullet(doc, "Completed emergency medical response training including triage, rapid assessment, and BSI/PPE protocols in live scenarios.")

add_section_heading(doc, "Work Experience")

add_entry_header(doc, "Team Lead", "Jan 2023 - Dec 2024")
add_org_line(doc, "Lift ATX | Austin, TX")
add_bullet(doc, "Led a team of 8 staff members, coordinating daily schedules and ensuring full shift coverage across a facility serving 2,000+ active members.")
add_bullet(doc, "Maintained safety, cleanliness, and supply readiness throughout the facility, responding to safety concerns and member requests promptly.")
add_bullet(doc, "Trained and oriented new team members on safety protocols, customer service standards, and daily operational procedures.")

add_entry_header(doc, "Project Analyst (Contract)", "Mar 2023 - Jun 2023")
add_org_line(doc, "TTM Analytics | Austin, TX")
add_bullet(doc, "Collected, entered, and analyzed demographic and health data to support community wellness programs in underserved populations.")

add_entry_header(doc, "Director of Operations and Communications", "Jun 2022 - Jan 2023")
add_org_line(doc, "Nuclear Baseball Training | Austin, TX")
add_bullet(doc, "Monitored athlete health metrics and enforced facility safety protocols; communicated updates to athletes and families.")

add_entry_header(doc, "National Scouting Coordinator", "Aug 2019 - May 2022")
add_org_line(doc, "Perfect Game USA | Austin, TX")
add_bullet(doc, "Collected and interpreted performance data for 200+ athletes annually with strict data entry accuracy and privacy standards.")

add_section_heading(doc, "Certifications")
add_cert_bullet(doc, "Basic Life Support (BLS/CPR/AED)", "American Heart Association, AHA Provider (2025)")
add_cert_bullet(doc, "EMT-Basic, Certified", "Texas Dept. of State Health Services (2025)")
add_cert_bullet(doc, "Basic Structure Fire Suppression", "Texas Commission on Fire Protection (2025)")

add_section_heading(doc, "Education")
add_edu_row(doc, "B.S. Sports Management, Public Relations", "Texas Tech University, Lubbock, TX (2019)")

doc.save('/Users/yahirgonzalez/Desktop/billy-cruz-resumes/Billy_Cruz_Resume_PCA.docx')
print("PCA .docx saved")


# ═══════════════════════════════════════════════
# BUILD TRANSPORTER RESUME
# ═══════════════════════════════════════════════
doc = setup_doc()

build_header(doc,
    "Billy James Cruz",
    "Transporter / Nursing Assistant Candidate | EMT-Certified | BLS Provider (AHA)",
    "Houston, TX Area", "(409) 673-4859", "billyjcruz@icloud.com")

add_section_heading(doc, "Professional Summary")
add_summary(doc, "BLS-certified EMT with direct patient transport experience on ambulance calls with the Seguin Fire Department. Trained in stretcher and wheelchair transport, oxygen tank handling, IV pump and Foley catheter management, infection control, and equipment disinfection. Over two years of customer service leadership at a 2,000-member facility. Conversational Spanish. Committed to Houston Methodist's ICARE values.")

add_section_heading(doc, "Skills")
add_skill_grid(doc, [
    "Patient Transportation and Safe Handling", "Basic Life Support (BLS/AHA)",
    "Stretcher and Wheelchair Operation", "Oxygen Tank Handling",
    "IV Pump and Foley Catheter Transport", "Equipment Disinfection",
    "Infection Control and Universal Precautions", "Two-Step Patient Identification",
    "Patient Safety Net (PSN) Reporting", "Customer Service",
    "Data Entry, Collection, and Interpretation", "Conversational Spanish",
])

add_section_heading(doc, "Healthcare Training and Clinical Experience")

add_entry_header(doc, "Emergency Medical Technician", "Jan 2025 - May 2025")
add_org_line(doc, "EMS University | San Antonio, TX | Clinical Rotations: Seguin Fire Department")
add_bullet(doc, "Transported inpatient and outpatient-level patients via stretcher and wheelchair during Seguin Fire Department clinical rotations.")
add_bullet(doc, "Managed IV bags, infusion pumps, and Foley catheters during transport; administered oxygen and verified cylinder storage.")
add_bullet(doc, "Verified patient identity via two-step protocol; delivered verbal handoff reports to receiving staff.")
add_bullet(doc, "Disinfected equipment before and after each use; tagged damaged items. Applied universal precautions and isolation procedures.")

add_entry_header(doc, "Recruit Fire Academy", "Aug 2024 - Oct 2024")
add_org_line(doc, "TEEX | College Station, TX")
add_bullet(doc, "Completed patient rescue and movement modules emphasizing safe lifting, body mechanics, and fall prevention.")

add_section_heading(doc, "Work Experience")

add_entry_header(doc, "Team Lead", "Jan 2023 - Dec 2024")
add_org_line(doc, "Lift ATX | Austin, TX")
add_bullet(doc, "Provided front-line customer service to 2,000+ active members, resolving requests promptly.")
add_bullet(doc, "Maintained facility safety: disinfected surfaces, inspected and tagged damaged equipment, kept walkways hazard-free.")
add_bullet(doc, "Managed 15+ daily tasks independently during 8-hour shifts in a high-volume, physically demanding environment.")
add_bullet(doc, "Trained and oriented 8+ new team members on safety protocols and daily workflow procedures.")

add_entry_header(doc, "Director of Operations and Communications", "Jun 2022 - Jan 2023")
add_org_line(doc, "Nuclear Baseball Training | Austin, TX")
add_bullet(doc, "Oversaw facility operations including safe equipment setup, maintenance, and storage; communicated health updates to athletes and families.")

add_entry_header(doc, "National Scouting Coordinator", "Aug 2019 - May 2022")
add_org_line(doc, "Perfect Game USA | Austin, TX")
add_bullet(doc, "Collected and interpreted performance data for 200+ athletes annually with strict data entry accuracy and privacy standards.")

add_section_heading(doc, "Certifications")
add_cert_bullet(doc, "Basic Life Support (BLS/CPR/AED)", "American Heart Association, AHA Provider (2025)")
add_cert_bullet(doc, "EMT-Basic, Certified", "Texas Dept. of State Health Services (2025)")
add_cert_bullet(doc, "Basic Structure Fire Suppression", "Texas Commission on Fire Protection (2025)")

add_section_heading(doc, "Education")
add_edu_row(doc, "B.S. Sports Management, Public Relations", "Texas Tech University, Lubbock, TX (2019)")

doc.save('/Users/yahirgonzalez/Desktop/billy-cruz-resumes/Billy_Cruz_Resume_Transporter.docx')
print("Transporter .docx saved")
